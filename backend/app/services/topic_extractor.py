from __future__ import annotations

import json
import math
import os
import re
from io import BytesIO
from typing import Any

import networkx as nx
import requests
from bs4 import BeautifulSoup
from docx import Document
from openai import OpenAI
from PyPDF2 import PdfReader


def get_openai_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not configured for topic extraction.")
    return OpenAI(api_key=api_key)


def extract_text_from_uploaded_file(filename: str, content: bytes) -> str:
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        reader = PdfReader(BytesIO(content))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    if lower_name.endswith(".docx"):
        doc = Document(BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)

    if lower_name.endswith((".txt", ".md", ".csv", ".json", ".html", ".htm")):
        return content.decode("utf-8", errors="ignore")

    # Best-effort fallback for unknown text-like uploads.
    return content.decode("utf-8", errors="ignore")


def fetch_full_content(source: str) -> str:
    if source.lower().endswith(".pdf") and os.path.exists(source):
        reader = PdfReader(source)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    if source.lower().endswith(".docx") and os.path.exists(source):
        doc = Document(source)
        return "\n".join(p.text for p in doc.paragraphs)

    if source.startswith("http://") or source.startswith("https://"):
        response = requests.get(source, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator=" ")
        return re.sub(r"\s+", " ", text).strip()

    if os.path.exists(source):
        with open(source, "r", encoding="utf-8", errors="ignore") as file:
            return file.read()

    return source


def chunk_text(text: str, max_chars: int = 8000) -> list[str]:
    chunks: list[str] = []
    start = 0
    length = len(text)

    while start < length:
        end = min(start + max_chars, length)
        if end < length:
            space_pos = text.rfind(" ", start, end)
            if space_pos != -1 and space_pos > start + int(max_chars * 0.6):
                end = space_pos
        chunks.append(text[start:end])
        start = end

    return chunks


def call_llm(prompt: str, temperature: float = 0.2) -> str:
    response = get_openai_client().chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
    )
    content = response.choices[0].message.content
    return content if content else ""


def parse_json_safely(text: str) -> Any:
    text = text.strip()

    try:
        return json.loads(text)
    except Exception:
        pass

    fenced = re.search(r"```(?:json)?(.*)```", text, re.DOTALL)
    if fenced:
        candidate = fenced.group(1).strip()
        try:
            return json.loads(candidate)
        except Exception:
            pass

    obj = re.search(r"(\{.*\})", text, re.DOTALL)
    if obj:
        candidate = obj.group(1).strip()
        try:
            return json.loads(candidate)
        except Exception:
            pass

    raise ValueError("Failed to parse JSON from LLM output")


def extract_topics_from_source(seminar_topic: str, source: str) -> list[str]:
    full_text = fetch_full_content(source)
    chunks = chunk_text(full_text, max_chars=8000)
    all_topics: list[str] = []

    for chunk in chunks:
        prompt = f"""
You are a research assistant for a seminar on: \"{seminar_topic}\".

From the text chunk below, extract all relevant technical topics for this seminar.
Focus on concepts, mechanisms, algorithms, architectures, and modern aspects.

Return JSON only in this format:
{{
  \"topics\": [\"Short Topic Name 1\", \"Short Topic Name 2\"]
}}

Text chunk:
\"\"\"{chunk}\"\"\"
"""

        try:
            out = call_llm(prompt)
            parsed = parse_json_safely(out)
            topics = parsed.get("topics", []) if isinstance(parsed, dict) else []
            if isinstance(topics, list):
                for topic in topics:
                    if isinstance(topic, str) and topic.strip():
                        all_topics.append(topic.strip())
        except Exception:
            continue

    seen: set[str] = set()
    unique_topics: list[str] = []
    for topic in all_topics:
        key = topic.lower()
        if key not in seen:
            seen.add(key)
            unique_topics.append(topic)

    return unique_topics


def get_embeddings(texts: list[str]) -> list[list[float]]:
    if not texts:
        return []

    response = get_openai_client().embeddings.create(model="text-embedding-3-small", input=texts)
    return [item.embedding for item in response.data]


def cosine_similarity(v1: list[float], v2: list[float]) -> float:
    if not v1 or not v2:
        return 0.0

    dot = sum(a * b for a, b in zip(v1, v2))
    norm1 = math.sqrt(sum(a * a for a in v1))
    norm2 = math.sqrt(sum(b * b for b in v2))

    if norm1 == 0 or norm2 == 0:
        return 0.0

    return dot / (norm1 * norm2)


def contextualize_topics_with_llm(
    seminar_topic: str,
    topics: list[str],
    batch_size: int = 40,
) -> list[str]:
    if not topics:
        return []

    refined_topics: list[str] = []

    for i in range(0, len(topics), batch_size):
        batch = topics[i : i + batch_size]
        prompt = f"""
You are refining topic names for a seminar on: \"{seminar_topic}\".

Rewrite each topic so it is specific to the seminar context when needed.

Rules:
1. Preserve the original meaning.
2. If the topic is too generic, add clear scope or context.
3. If the topic is already specific, keep it almost unchanged.
4. Keep each topic concise (3-12 words).
5. Return the same number of topics and in the same order.

Return JSON only in this exact format:
{{
  \"topics\": [\"Refined Topic 1\", \"Refined Topic 2\"]
}}

Input topics:
{json.dumps(batch, ensure_ascii=False, indent=2)}
"""

        try:
            out = call_llm(prompt, temperature=0.1)
            parsed = parse_json_safely(out)
            candidate = parsed.get("topics", []) if isinstance(parsed, dict) else []
            if isinstance(candidate, list) and len(candidate) == len(batch):
                for original, rewritten in zip(batch, candidate):
                    if isinstance(rewritten, str) and rewritten.strip():
                        refined_topics.append(rewritten.strip())
                    else:
                        refined_topics.append(original)
            else:
                refined_topics.extend(batch)
        except Exception:
            refined_topics.extend(batch)

    return refined_topics


def cluster_topics_with_lpa(
    topics: list[str],
    similarity_threshold: float = 0.65,
) -> list[list[str]]:
    if not topics:
        return []

    embeddings = get_embeddings(topics)

    graph = nx.Graph()
    for idx, topic in enumerate(topics):
        graph.add_node(idx, label=topic)

    topic_count = len(topics)
    for i in range(topic_count):
        for j in range(i + 1, topic_count):
            sim = cosine_similarity(embeddings[i], embeddings[j])
            if sim >= similarity_threshold:
                graph.add_edge(i, j, weight=sim)

    communities = list(nx.algorithms.community.label_propagation_communities(graph))

    clusters: list[list[str]] = []
    for community in communities:
        cluster = [graph.nodes[idx]["label"] for idx in community]
        clusters.append(sorted(cluster))

    clusters.sort(key=len, reverse=True)
    return clusters


def topic_extraction_agent(
    seminar_topic: str,
    sources: list[str],
    similarity_threshold: float = 0.65,
) -> dict[str, Any]:
    all_topics: list[str] = []

    for source in sources:
        all_topics.extend(extract_topics_from_source(seminar_topic, source))

    seen: set[str] = set()
    merged_topics: list[str] = []
    for topic in all_topics:
        key = topic.lower()
        if key not in seen:
            seen.add(key)
            merged_topics.append(topic)

    refined_topics = contextualize_topics_with_llm(seminar_topic, merged_topics)

    seen_refined: set[str] = set()
    merged_refined_topics: list[str] = []
    for topic in refined_topics:
        key = topic.lower()
        if key not in seen_refined:
            seen_refined.add(key)
            merged_refined_topics.append(topic)

    clusters = cluster_topics_with_lpa(
        merged_refined_topics,
        similarity_threshold=similarity_threshold,
    )

    return {
        "all_topics": merged_refined_topics,
        "clusters": clusters,
    }


def summarize_clusters_with_llm(seminar_topic: str, clusters: list[list[str]]) -> str:
    prompt = f"""
You are helping organize topics for a seminar on: \"{seminar_topic}\".

You are given clusters of related technical topics.
For each cluster:
1. Give a short, domain-specific name.
2. Provide a one-sentence description.
3. List member topics as bullet points.

Return markdown only.

Clusters:
{json.dumps(clusters, ensure_ascii=False, indent=2)}
"""
    return call_llm(prompt)
